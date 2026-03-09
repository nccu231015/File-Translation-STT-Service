"""
transcript_docx_service.py

Generates a bilingual transcript Word document from Whisper segments.

Each entry in the document looks like:
  [00:00:01 → 00:00:05]
  原文：大家好，今天我們來討論...
  譯文：Hello everyone, today we discuss...
"""

from io import BytesIO
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _fmt_time(seconds: float) -> str:
    """Format seconds to HH:MM:SS."""
    td = timedelta(seconds=int(seconds))
    total_seconds = int(td.total_seconds())
    h = total_seconds // 3600
    m = (total_seconds % 3600) // 60
    s = total_seconds % 60
    return f"{h:02d}:{m:02d}:{s:02d}"


class TranscriptDocxService:
    """
    Generates a bilingual (original + translation) Word document
    from a list of translated segments.

    Each segment dict expected:
      { start: float, end: float, original: str, translated: str }
    """

    def generate(
        self,
        file_name: str,
        segments: list[dict],
        src_lang: str,
        tgt_lang: str,
        date: datetime = None,
    ) -> bytes:
        """
        Args:
            file_name   : Original audio file name (used in title)
            segments    : List of {start, end, original, translated} dicts
            src_lang    : Source language label (e.g. "中文")
            tgt_lang    : Target language label (e.g. "英文")
            date        : Meeting date (default: now)

        Returns:
            bytes of the generated .docx file
        """
        doc = Document()

        # ── Title ─────────────────────────────────────────────────────────
        title_para = doc.add_heading("雙語逐字稿", 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ── Metadata ──────────────────────────────────────────────────────
        info = doc.add_paragraph()
        info.add_run("音訊檔案：").bold = True
        info.add_run(file_name)

        info2 = doc.add_paragraph()
        info2.add_run("產生時間：").bold = True
        info2.add_run((date or datetime.now()).strftime("%Y年%m月%d日 %H:%M"))

        info3 = doc.add_paragraph()
        info3.add_run("語言對照：").bold = True
        info3.add_run(f"{src_lang}  ↔  {tgt_lang}")

        doc.add_paragraph()  # spacing
        doc.add_paragraph("─" * 60)
        doc.add_paragraph()

        # ── Segments ──────────────────────────────────────────────────────
        for seg in segments:
            start_str = _fmt_time(seg.get("start", 0))
            end_str = _fmt_time(seg.get("end", 0))
            original = seg.get("original", "").strip()
            translated = seg.get("translated", "").strip()

            # Time stamp header
            ts_para = doc.add_paragraph()
            ts_run = ts_para.add_run(f"[{start_str} → {end_str}]")
            ts_run.bold = True
            ts_run.font.color.rgb = RGBColor(90, 90, 90)
            ts_run.font.size = Pt(9)

            # Original line
            orig_para = doc.add_paragraph()
            label_run = orig_para.add_run(f"原文：")
            label_run.bold = True
            label_run.font.color.rgb = RGBColor(30, 100, 200)
            orig_run = orig_para.add_run(original)
            orig_run.font.color.rgb = RGBColor(30, 100, 200)

            # Translation line
            trans_para = doc.add_paragraph()
            t_label = trans_para.add_run(f"譯文：")
            t_label.bold = True
            t_label.font.color.rgb = RGBColor(34, 139, 34)
            t_run = trans_para.add_run(translated if translated else "(翻譯失敗)")
            t_run.font.color.rgb = RGBColor(34, 139, 34)

            # Separator
            doc.add_paragraph()

        # ── Footer note ───────────────────────────────────────────────────
        doc.add_paragraph("─" * 60)
        note = doc.add_paragraph()
        note_run = note.add_run("本逐字稿由 AI 自動生成，如有誤差請以實際錄音為準。")
        note_run.font.size = Pt(9)
        note_run.font.color.rgb = RGBColor(150, 150, 150)

        # ── Serialize ─────────────────────────────────────────────────────
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()
