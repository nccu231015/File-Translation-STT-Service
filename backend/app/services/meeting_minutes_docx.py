"""
meeting_minutes_docx.py

Generates a bilingual (Traditional Chinese / English) meeting minutes Word document.
Each section displays the Chinese content first, followed by its English translation.
"""

import os
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime


class MeetingMinutesDocxService:
    """
    Service for generating professional, bilingual Word documents
    from meeting analysis results.
    """

    def generate_minutes(
        self,
        file_name: str,
        meeting_objective: str = "",
        discussion_summary=None,
        decisions: list = None,
        action_items: list = None,
        attendees=None,
        schedule_notes: str = "",
        date: datetime = None,
        # English counterparts (optional — if provided, show bilingual)
        en_meeting_objective: str = "",
        en_discussion_summary=None,
        en_decisions: list = None,
        en_action_items: list = None,
        en_schedule_notes: str = "",
    ) -> bytes:
        """
        Generate a bilingual meeting minutes Word document.

        Args:
            file_name           : Original audio file name
            *                   : Chinese content fields
            en_*                : English translations of each field (optional)
            date                : Meeting date (default: now)

        Returns:
            Bytes of the generated .docx file
        """
        doc = Document()

        decisions = decisions or []
        action_items = action_items or []
        attendees = attendees or []
        en_decisions = en_decisions or []
        en_action_items = en_action_items or []

        # ── Helpers ───────────────────────────────────────────────────────────
        def format_content(val):
            """Recursively stringify nested structures."""
            if isinstance(val, list):
                if val and isinstance(val[0], dict):
                    text = ""
                    for item in val:
                        for k, v in item.items():
                            key_display = k.replace('_', ' ').capitalize()
                            text += f"{key_display}: {format_content(v)}\n"
                        text += "\n"
                    return text.strip()
                return "\n".join([str(v) for v in val])
            elif isinstance(val, dict):
                return "\n".join(f"{k}: {v}" for k, v in val.items())
            return str(val) if val is not None else ""

        def add_bilingual_para(doc, zh_text: str, en_text: str):
            """Add a paragraph with Chinese text above and indented English below."""
            if zh_text:
                p_zh = doc.add_paragraph()
                run_zh = p_zh.add_run(zh_text)
                run_zh.font.color.rgb = RGBColor(30, 80, 180)

            if en_text:
                p_en = doc.add_paragraph()
                p_en.paragraph_format.left_indent = Inches(0.3)
                run_label = p_en.add_run("EN: ")
                run_label.bold = True
                run_label.font.color.rgb = RGBColor(100, 100, 100)
                run_label.font.size = Pt(9)
                run_en = p_en.add_run(en_text)
                run_en.font.color.rgb = RGBColor(100, 100, 100)
                run_en.font.size = Pt(9)

        def add_lang_label(para, label: str, color: RGBColor):
            run = para.add_run(label)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = color
            return run

        # ── Title ─────────────────────────────────────────────────────────────
        title = doc.add_heading('會議記錄 Meeting Minutes', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ── 一、基本資訊 ──────────────────────────────────────────────────────
        doc.add_heading('一、會議基本資訊  |  1. Meeting Information', level=1)

        p = doc.add_paragraph()
        p.add_run('會議主題 Topic：').bold = True
        p.add_run(file_name or '未命名會議')

        p = doc.add_paragraph()
        p.add_run('會議日期 Date：').bold = True
        p.add_run((date or datetime.now()).strftime('%Y年%m月%d日 %H:%M'))

        p = doc.add_paragraph()
        p.add_run('會議方式 Mode：').bold = True
        p.add_run('線上會議 / Online Meeting')

        # Attendees
        if attendees:
            if isinstance(attendees, str):
                if ',' in attendees or '、' in attendees or '；' in attendees:
                    attendees = [a.strip() for a in attendees.replace('、', ',').replace('；', ',').split(',') if a.strip()]
                else:
                    attendees = [attendees]

            p = doc.add_paragraph()
            p.add_run('與會人員 Attendees：').bold = True
            doc.add_paragraph()
            for attendee in attendees:
                if attendee and str(attendee).strip():
                    doc.add_paragraph(str(attendee).strip(), style='List Bullet')

        doc.add_paragraph()

        # ── 二、會議目的 ──────────────────────────────────────────────────────
        doc.add_heading('二、會議目的  |  2. Meeting Objective', level=1)

        zh_obj = format_content(meeting_objective) if meeting_objective else '依會議內容進行需求討論與確認'
        en_obj = format_content(en_meeting_objective) if en_meeting_objective else ''
        add_bilingual_para(doc, zh_obj, en_obj)
        doc.add_paragraph()

        # ── 三、討論摘要 ──────────────────────────────────────────────────────
        doc.add_heading('三、討論摘要  |  3. Discussion Summary', level=1)

        def split_paragraphs(val):
            text = format_content(val) if val else ''
            return [p.strip() for p in text.split('\n\n') if p.strip()]

        zh_paras = split_paragraphs(discussion_summary)
        en_paras = split_paragraphs(en_discussion_summary)

        max_paras = max(len(zh_paras), len(en_paras), 1)
        for i in range(max_paras):
            zh_p = zh_paras[i] if i < len(zh_paras) else ''
            en_p = en_paras[i] if i < len(en_paras) else ''
            if zh_p or en_p:
                add_bilingual_para(doc, zh_p, en_p)
                doc.add_paragraph()

        # ── 四、決策事項 ──────────────────────────────────────────────────────
        doc.add_heading('四、決策事項  |  4. Decisions', level=1)

        if decisions:
            for idx, decision in enumerate(decisions, 1):
                en_decision = en_decisions[idx - 1] if idx - 1 < len(en_decisions) else ''

                # Chinese line
                p_zh = doc.add_paragraph()
                num_run = p_zh.add_run(f'{idx}. ')
                num_run.bold = True
                num_run.font.color.rgb = RGBColor(34, 120, 34)
                txt_run = p_zh.add_run(str(decision))
                txt_run.font.color.rgb = RGBColor(34, 120, 34)

                # English line (indented)
                if en_decision:
                    p_en = doc.add_paragraph()
                    p_en.paragraph_format.left_indent = Inches(0.3)
                    add_lang_label(p_en, "EN: ", RGBColor(100, 130, 100))
                    en_run = p_en.add_run(str(en_decision))
                    en_run.font.color.rgb = RGBColor(100, 130, 100)
                    en_run.font.size = Pt(9)

                doc.add_paragraph()
        else:
            doc.add_paragraph('本次會議無決策事項  |  No decisions made in this meeting.')

        doc.add_paragraph()

        # ── 五、時程安排 ──────────────────────────────────────────────────────
        doc.add_heading('五、時程與後續安排  |  5. Schedule & Follow-Up', level=1)

        zh_schedule = format_content(schedule_notes) if schedule_notes else '依會議討論結果進行後續規劃與執行'
        en_schedule = format_content(en_schedule_notes) if en_schedule_notes else ''
        add_bilingual_para(doc, zh_schedule, en_schedule)
        doc.add_paragraph()

        # ── 六、待辦事項 ──────────────────────────────────────────────────────
        doc.add_heading('六、待辦事項  |  6. Action Items', level=1)

        if action_items:
            for idx, action in enumerate(action_items):
                en_action = en_action_items[idx] if idx < len(en_action_items) else {}

                # Format Chinese action item
                if isinstance(action, dict):
                    task = str(action.get('task', ''))
                    owner = str(action.get('owner', ''))
                    deadline = str(action.get('deadline', ''))
                    zh_text = f"【{owner}】{task}" if owner else task
                    if deadline and deadline.lower() not in ['ongoing', '持續進行']:
                        zh_text += f" (期限: {deadline})"
                else:
                    zh_text = str(action)

                # Format English action item
                en_text = ''
                if isinstance(en_action, dict):
                    en_task = str(en_action.get('task', ''))
                    en_owner = str(en_action.get('owner', ''))
                    en_deadline = str(en_action.get('deadline', ''))
                    en_text = f"[{en_owner}] {en_task}" if en_owner else en_task
                    if en_deadline and en_deadline.lower() != 'ongoing':
                        en_text += f" (Due: {en_deadline})"
                elif en_action:
                    en_text = str(en_action)

                p_zh = doc.add_paragraph(zh_text, style='List Bullet')
                for run in p_zh.runs:
                    run.font.color.rgb = RGBColor(30, 100, 200)

                if en_text:
                    p_en = doc.add_paragraph()
                    p_en.paragraph_format.left_indent = Inches(0.5)
                    add_lang_label(p_en, "EN: ", RGBColor(100, 130, 170))
                    en_run = p_en.add_run(en_text)
                    en_run.font.color.rgb = RGBColor(100, 130, 170)
                    en_run.font.size = Pt(9)

                doc.add_paragraph()
        else:
            doc.add_paragraph('本次會議無待辦事項  |  No action items.')

        doc.add_paragraph()

        # ── 七、備註 ──────────────────────────────────────────────────────────
        doc.add_heading('七、備註  |  7. Notes', level=1)
        note = doc.add_paragraph(
            '本會議記錄由語音轉文字系統自動生成，如有遺漏或錯誤，請以實際會議內容為準。\n'
            'This meeting record was automatically generated by an AI speech-to-text system. '
            'Please verify against the actual meeting content.'
        )
        for run in note.runs:
            run.font.size = Pt(9)

        # ── Serialize ─────────────────────────────────────────────────────────
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()
