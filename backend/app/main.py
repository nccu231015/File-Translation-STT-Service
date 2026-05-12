import os
import re
import shutil
import tempfile
import httpx
import json
import base64
import datetime
from fastapi import FastAPI, UploadFile, File, HTTPException, Body, Response, Form, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from app.services.stt_service import stt_service
from app.services.llm_service import llm_service
from app.services.pdf_service import pdf_service
from app.services.gpu_live_monitor import (
    MODULE_KM_ASSISTANT,
    MODULE_VOICE,
    gpu_work_acquire,
    gpu_work_release,
    gpu_work_scope,
    get_active_modules,
)
from app.services.meeting_minutes_docx import MeetingMinutesDocxService
from app.services.transcript_docx_service import TranscriptDocxService
from app.services.employee_db import (
    get_employee, get_subordinates,
    save_employee_record, get_employee_records,
    ensure_records_table,
)
from app.services.rank_service import get_rank, has_view_permission
from app.services.factory.factory_agent import FactoryAgentService
from app.services.factory.factory_redis import factory_store
from app.services.factory.doc_redis import doc_store
from pydantic import BaseModel
from typing import Optional
import uuid
import traceback
import asyncio
from fastapi.concurrency import run_in_threadpool
from pdf2docx import Converter
from app.services.factory.sql_tools import FactorySqlTools
import docx
from docx.shared import Pt
from docx.oxml.ns import qn
import subprocess
import platform
try:
    from docx2pdf import convert as docx2pdf_convert
except ImportError:
    docx2pdf_convert = None

app = FastAPI()
factory_agent = FactoryAgentService(llm_service)

@app.on_event("startup")
async def startup_event():
    """Initialization service: ensure DB connection and tables are ready"""
    print("\n[Startup] Application is starting...", flush=True)
    
    # 1. Employee DB (MySQL)
    try:
        await ensure_records_table()
        print("[EmployeeDB] ✓ employee_records table ready")
    except Exception as e:
        print(f"[EmployeeDB Error] {e}")

    # 2. Factory Session Store (Redis/Memory)
    try:
        await factory_store.connect()
        print("[FactoryStore] ✓ Session store initialized")
    except Exception as e:
        print(f"[FactoryStore Error] {e}")

    # 2b. Document KM Session Store (same Redis, different key prefix)
    try:
        await doc_store.connect()
        print("[DocStore] ✓ Document session store initialized")
    except Exception as e:
        print(f"[DocStore Error] {e}")

    # 3. Factory Databases Health Check
    try:
        factory_tools = FactorySqlTools()
        db_status = factory_tools.test_connections()
        print(f"[HealthCheck] MSSQL: {db_status['mssql']}, Postgres: {db_status['postgres']}")
    except Exception as e:
        print(f"[HealthCheck Error] {e}")

# Middleware
@app.middleware("http")
async def add_ngrok_header(request, call_next):
    response = await call_next(request)
    response.headers["ngrok-skip-browser-warning"] = "true"
    return response

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ── System Status ─────────────────────────────────────────────────────────────
@app.get("/system-status")
async def system_status():
    """Return a lightweight system load snapshot for the frontend.

    - active_modules: which GPU-heavy modules are currently running and how many tasks.
    - translate_slots_used: how many PDF translation parallel slots are currently occupied.
    - translate_slots_total: the configured max parallel slots for translation.
    - busy: True when translation is near full capacity (>= 80% of slots used).
    """
    active = get_active_modules()
    translate_total = llm_service._translate_batch_sem._value + (
        # _value decreases as slots are acquired; compute used = total - available
        # asyncio.Semaphore stores current available count in _value
        0  # placeholder; real computation below
    )
    sem = llm_service._translate_batch_sem
    # asyncio.Semaphore._value = remaining available slots
    translate_available = sem._value
    translate_total_cfg = int(os.getenv("OLLAMA_TRANSLATE_MAX_PARALLEL", "15"))
    translate_used = max(0, translate_total_cfg - translate_available)

    busy = translate_used >= int(translate_total_cfg * 0.8)

    return {
        "busy": busy,
        "active_modules": active,
        "translate_slots_used": translate_used,
        "translate_slots_total": translate_total_cfg,
    }


class LoginRequest(BaseModel):
    username: str
    password: str

@app.post("/api/login")
async def login(body: LoginRequest):
    LDAP_URL = "http://api.jebsee.com.tw:8081/api/LDAPLogin"
    try:
        async with httpx.AsyncClient(timeout=10.0) as client:
            resp = await client.post(LDAP_URL, data={"username": body.username, "password": body.password})
        data = resp.json()
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"LDAP service unavailable: {e}")

    if str(data.get("chkIdentity", "false")).lower() != "true":
        raise HTTPException(status_code=401, detail=data.get("msg", "Login failed"))

    sso_username = data.get("username", body.username)
    emp = await get_employee(sso_username)
    if emp:
        title, dept, name = emp.get("DUTYNAME", ""), emp.get("DEPTNAME", data.get("dpt", "")), emp.get("EMPNAME", data.get("name", ""))
    else:
        title, dept, name = "", data.get("dpt", ""), data.get("name", "")

    rank = get_rank(title)
    can_view = has_view_permission(rank)
    return {"username": sso_username, "name": name, "dpt": dept, "title": title, "rank": rank, "canViewRecords": can_view}

@app.get("/api/records/{empid}")
async def get_records(empid: str):
    """
    Get employee records list.
    Manager rank: gets subordinates list; Employee rank: gets personal history.
    """
    requester = await get_employee(empid)
    if not requester:
        # Try to get individual records directly
        records = await get_employee_records(empid)
        return {"empid": empid, "employees": [], "records": records}
    
    rank = get_rank(requester.get("DUTYNAME", ""))
    if has_view_permission(rank):
        # Manager view: get subordinates list
        dept = requester.get("DEPTNAME", "")
        subordinates = await get_subordinates(empid, dept, rank)
        return {"requester_rank": rank, "employees": subordinates}
    else:
        # Personal view: get personal records
        records = await get_employee_records(empid)
        return {"empid": empid, "records": records}

@app.post("/api/employee-records")
async def create_employee_record(body: dict):
    # Compatible with old and new Request Body versions
    record_id = await save_employee_record(
        empid=body.get("empid"),
        record_type=body.get("type", "voice"),
        file_name=body.get("file_name", "unknown"),
        summary=body.get("summary", ""),
        decisions=body.get("decisions", ""),
        action_items=body.get("action_items", ""),
    )
    if record_id is None: raise HTTPException(status_code=500, detail="Failed to save record")
    return {"id": record_id, "success": True}

@app.get("/")
async def read_root():
    return {"message": "Factory AI & STT Service is Running"}

@app.post("/stt")
async def transcribe_audio(file: UploadFile = File(...), mode: str = Form("chat")):
    """
    Legacy STT route, now simplified for basic transcription and chat assistants.
    Advanced meeting analysis and parameter tuning tasks are handled by /api/v1/stt/process.
    """
    temp_file_path = ""
    try:
        suffix = os.path.splitext(file.filename)[1]
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp_file:
            shutil.copyfileobj(file.file, temp_file)
            temp_file_path = temp_file.name
        
        # Perform Whisper transcription (parallelized via threadpool)
        stt_result = await run_in_threadpool(stt_service.transcribe, temp_file_path)
        if os.path.exists(temp_file_path): os.remove(temp_file_path)
        
        user_text = stt_result["text"]
        
        # Keep only basic chat for assistant features
        if mode == "chat":
            llm_response = await llm_service.chat(user_text)
            return {"transcription": stt_result, "llm_response": llm_response}
        
        return {"transcription": stt_result}
    except Exception as e:
        if temp_file_path and os.path.exists(temp_file_path): os.remove(temp_file_path)
        raise HTTPException(status_code=500, detail=str(e))

# ------------------------------------------------------------------------------
# n8n Microservice Interface: POST /api/v1/stt/process
# Designed specifically for n8n Webhook -> HTTP Request -> Respond to Webhook flows.
# Independent from legacy /stt, supporting dynamic LLM parameter tuning.
# Supports both stt_only mode and full meeting minutes generation with Word exports.
# ------------------------------------------------------------------------------
@app.post("/api/v1/stt/process")
async def stt_process_for_n8n(
    file: UploadFile = File(...),
    mode: str = Form("stt_only"),          # "stt_only" or "minutes"
    temperature: float = Form(0.2),        # LLM temperature (0.0 strict ~ 1.0 creative)
    num_predict: int = Form(1024),         # Max tokens for LLM response
    num_ctx: int = Form(4096),             # Context window size (override from n8n)
    model: str = Form(""),                 # Ollama model override (empty = use service default)
    initial_prompt: str = Form(""),        # Whisper initial prompt for terminology/style
    language: str = Form("zh"),            # Designated language for Whisper
):
    """
    n8n STT Microservice Entrypoint.
    - Receives audio + tunable LLM/Whisper parameters.
    - Uses run_in_threadpool for non-blocking parallel processing.
    - Returns JSON for n8n's 'Respond to Webhook' node to forward results.
    - In 'minutes' mode, generates bilingual Word documents (meeting minutes + transcript).
    """
    temp_file_path = ""
    try:
        # Step 1: Save uploaded audio to temp location
        suffix = os.path.splitext(file.filename or "audio")[1] or ".wav"
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            shutil.copyfileobj(file.file, tmp)
            temp_file_path = tmp.name
        print(f"[n8n STT] Received: {file.filename} | mode={mode} | temp={temp_file_path}", flush=True)
        try:
            _nbytes = os.path.getsize(temp_file_path)
        except OSError:
            _nbytes = -1
        print(f"[n8n STT] temp audio bytes_on_disk={_nbytes}", flush=True)

        if mode == "minutes":
            gpu_work_acquire(MODULE_VOICE)
        try:
            # Step 2: Whisper Transcription (using threadpool for concurrency)
            stt_result = await run_in_threadpool(stt_service.transcribe, temp_file_path)
            transcript_text = stt_result.get("text", "")
            print(f"[n8n STT] Transcription done: {len(transcript_text)} chars", flush=True)

            # Clean up temp audio file
            if os.path.exists(temp_file_path):
                os.remove(temp_file_path)
                temp_file_path = ""

            # Step 3: Return raw transcript if mode is stt_only
            if mode == "stt_only":
                return {
                    "status": "success",
                    "mode": "stt_only",
                    "transcript": transcript_text,
                    "language": stt_result.get("language", ""),
                    "processing_time": stt_result.get("processing_time", 0),
                    "segments": stt_result.get("segments", []),
                }

            # ── minutes mode: Full analysis + bilingual Word document generation ────

            # Step 4a: LLM analysis (Traditional Chinese output)
            llm_options = {"temperature": temperature, "num_predict": num_predict, "num_ctx": num_ctx}
            model_override = model.strip() or None
            print(f"[n8n STT] Running LLM analysis | model={model_override or 'default'} options={llm_options}", flush=True)
            analysis = await run_in_threadpool(
                llm_service.analyze_meeting_transcript, transcript_text,
                model_override, temperature, num_predict, num_ctx
            )

            # Step 4b: Validity check — if key fields are all empty, the LLM JSON parse
            # likely failed and triggered fallback. Retry once with a larger num_predict
            # to give the model more room to output complete JSON.
            _is_analysis_empty = (
                not analysis.get("meeting_objective", "").strip()
                and not analysis.get("decisions", [])
                and not analysis.get("action_items", [])
                and not analysis.get("attendees", "")
            )
            if _is_analysis_empty:
                _retry_num_predict = max(num_predict * 2, 2048)
                print(
                    f"[n8n STT] Analysis key fields are all empty (likely JSON parse fallback). "
                    f"Retrying with num_predict={_retry_num_predict} ...",
                    flush=True,
                )
                analysis = await run_in_threadpool(
                    llm_service.analyze_meeting_transcript, transcript_text,
                    model_override, temperature, _retry_num_predict
                )
                print(
                    f"[n8n STT] Retry analysis complete. "
                    f"meeting_objective={bool(analysis.get('meeting_objective'))} "
                    f"decisions={len(analysis.get('decisions', []))} "
                    f"action_items={len(analysis.get('action_items', []))}",
                    flush=True,
                )

            # Step 5: Translate analysis to English (for bilingual minutes)
            try:
                en_analysis = await run_in_threadpool(llm_service.translate_analysis, analysis, model_override)
            except Exception as te:
                print(f"[n8n STT] translate_analysis failed: {te}")
                en_analysis = {}

            # Step 6: Generate meeting minutes DOCX
            file_download = None
            try:
                minutes_svc = MeetingMinutesDocxService()
                minutes_bytes = await run_in_threadpool(
                    minutes_svc.generate_minutes,
                    file.filename,
                    analysis.get("meeting_objective", ""),
                    analysis.get("discussion_summary", ""),
                    analysis.get("decisions", []),
                    analysis.get("action_items", []),
                    analysis.get("attendees", []),
                    analysis.get("schedule_notes", ""),
                    None,
                    en_analysis.get("meeting_objective", ""),
                    en_analysis.get("discussion_summary", ""),
                    en_analysis.get("decisions", []),
                    en_analysis.get("action_items", []),
                    en_analysis.get("schedule_notes", ""),
                )
                file_download = {
                    "filename": f"meeting_minutes_{file.filename}.docx",
                    "content_base64": base64.b64encode(minutes_bytes).decode("utf-8"),
                    "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                }
            except Exception as de:
                print(f"[n8n STT] MeetingMinutesDocx generation failed: {de}")

            # Step 7: Translate segments and generate bilingual transcript DOCX
            translated_segments = []
            transcript_download = None
            segments = stt_result.get("segments", [])
            detected_lang = stt_result.get("language", "zh")
            is_chinese = detected_lang.lower().startswith("zh")
            try:
                if segments:
                    translated_segments = await llm_service.translate_segments_async(
                        segments, detected_lang, model_override
                    )
                    transcript_svc = TranscriptDocxService()
                    transcript_bytes = await run_in_threadpool(
                        transcript_svc.generate,
                        file.filename,
                        translated_segments,
                        "Chinese" if is_chinese else "English",
                        "English" if is_chinese else "Chinese (Traditional)",
                    )
                    transcript_download = {
                        "filename": f"transcript_{file.filename}.docx",
                        "content_base64": base64.b64encode(transcript_bytes).decode("utf-8"),
                        "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    }
            except Exception as te2:
                print(f"[n8n STT] TranscriptDocx generation failed: {te2}")

            # Step 8: Build bilingual composite display fields for frontend
            composite_summary = ""
            zh_obj = analysis.get("meeting_objective", "").strip()
            en_obj = en_analysis.get("meeting_objective", "").strip()
            if zh_obj or en_obj:
                composite_summary += f"【會議目的 | Meeting Objective】\n{zh_obj}\n{en_obj}\n\n"

            zh_sum = analysis.get("discussion_summary", analysis.get("summary", "")).strip()
            en_sum = en_analysis.get("discussion_summary", en_analysis.get("summary", "")).strip()
            if zh_sum or en_sum:
                composite_summary += f"【討論摘要 | Discussion Summary】\n{zh_sum}\n{en_sum}"

            zh_decisions = analysis.get("decisions", [])
            en_decisions = en_analysis.get("decisions", [])
            composite_decisions = []
            for i in range(max(len(zh_decisions), len(en_decisions))):
                zh_d = zh_decisions[i] if i < len(zh_decisions) else ""
                en_d = en_decisions[i] if i < len(en_decisions) else ""
                composite_decisions.append(f"{zh_d}\n{en_d}".strip())

            zh_actions = analysis.get("action_items", [])
            en_actions = en_analysis.get("action_items", [])
            composite_actions = []
            for i in range(max(len(zh_actions), len(en_actions))):
                zh_a = zh_actions[i] if i < len(zh_actions) else {}
                en_a = en_actions[i] if i < len(en_actions) else {}
                if isinstance(zh_a, str):
                    en_task = en_a if isinstance(en_a, str) else en_a.get("task", "")
                    composite_actions.append(f"{zh_a}\n{en_task}".strip())
                else:
                    task = f"{zh_a.get('task', '')}\n{en_a.get('task', '')}".strip()
                    composite_actions.append(
                        {"task": task, "owner": zh_a.get("owner", ""), "deadline": zh_a.get("deadline", "")}
                    )

            # Step 9: Construct final response
            result = {
                "status": "success",
                "mode": "minutes",
                "transcript": transcript_text,
                "language": stt_result.get("language", ""),
                "processing_time": stt_result.get("processing_time", 0),
                "summary": composite_summary.strip() or zh_sum,
                "meeting_objective": analysis.get("meeting_objective", ""),
                "decisions": composite_decisions,
                "action_items": composite_actions,
                "attendees": analysis.get("attendees", []),
                "llm_options_used": llm_options,
            }
            if translated_segments:
                result["translated_segments"] = translated_segments
            if file_download:
                result["file_download"] = file_download
            if transcript_download:
                result["transcript_download"] = transcript_download
            return result
        finally:
            if mode == "minutes":
                gpu_work_release(MODULE_VOICE)

    except Exception as e:
        if temp_file_path and os.path.exists(temp_file_path):
            os.remove(temp_file_path)
        print(f"[n8n STT Error] {e}\n{traceback.format_exc()}", flush=True)
        raise HTTPException(status_code=500, detail=f"STT processing failed: {str(e)}")

@app.post("/pdf-translation")
async def translate_pdf(background_tasks: BackgroundTasks, file: UploadFile = File(...), target_lang: str = Form(None), debug: str = Form("false"), is_complex_table: str = Form("false"), num_ctx: int = Form(None)):
    debug_mode = str(debug).lower() in ("true", "1", "t", "yes", "on")
    is_complex_table_bool = str(is_complex_table).lower() in ("true", "1", "t", "yes", "on")
    temp_input_path = ""
    docx_path = ""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_input:
            shutil.copyfileobj(file.file, temp_input)
            temp_input_path = temp_input.name
        
        try:
            _psz = os.path.getsize(temp_input_path)
        except OSError:
            _psz = -1
        print(
            f"[PDF] translate_pdf temp_saved bytes={_psz} debug={debug_mode} complex_table={is_complex_table_bool}",
            flush=True,
        )
        result_list = await pdf_service.process_pdf(temp_input_path, force_target_lang=target_lang, debug_mode=debug_mode, is_complex_table=is_complex_table_bool, num_ctx=num_ctx)
        output_pdf_path = result_list[0]["file_path"]
        
        with open(output_pdf_path, 'rb') as f:
            pdf_b64 = base64.b64encode(f.read()).decode('utf-8')
            
        # Docx generation and Table translation
        docx_b64 = None
        if not debug_mode:
            try:
                docx_path = output_pdf_path.replace(".pdf", ".docx")
                
                def _convert():
                    cv = Converter(output_pdf_path)
                    cv.convert(docx_path)
                    cv.close()
                    
                await run_in_threadpool(_convert)

                # ---- Stage 2: Process tables in Word (DOCX) ----
                def _process_docx_tables():
                    doc = docx.Document(docx_path)
                    to_translate = {}
                    
                    def collect_text(container):
                        def _has_chinese(text: str) -> bool:
                            return any('\u4e00' <= c <= '\u9fff' or '\u3400' <= c <= '\u4dbf'
                                       or '\uf900' <= c <= '\ufaff' for c in text)
                        
                        _tgt = (target_lang or "").lower()
                        _target_is_chinese = any(code in _tgt for code in ['zh', 'tw', 'cn', 'hk'])

                        if hasattr(container, 'tables'):
                            for t_idx, table in enumerate(container.tables):
                                for r_idx, row in enumerate(table.rows):
                                    for c_idx, cell in enumerate(row.cells):
                                        key = f"t{t_idx}_r{r_idx}_c{c_idx}"
                                        
                                        # Use full-cell collection to prevent fragmenting technical names 
                                        # (Reference style logic from before)
                                        if hasattr(cell, 'paragraphs'):
                                            full_text = "".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                                            
                                            should_translate = False
                                            if _target_is_chinese:
                                                # To Chinese -> pick up Latin only if it's not mostly numbers
                                                if re.search(r'[a-zA-Z]{2,}', full_text) and not _has_chinese(full_text):
                                                    should_translate = True
                                            else:
                                                # To English -> pick up only if Chinese is present
                                                if _has_chinese(full_text):
                                                    should_translate = True
                                            
                                            if should_translate:
                                                to_translate[key] = full_text
                                        collect_text(cell)
                    
                    to_translate = {} # Changed to dict to keep track of keys for faster replacement
                    collect_text(doc)
                    cell_texts = to_translate  # key -> full_text mapping
                    return doc, list(set(to_translate.values())), cell_texts
                
                if is_complex_table_bool:
                    doc, content_list, cell_texts = await run_in_threadpool(_process_docx_tables)
                else:
                    doc, content_list, cell_texts = None, [], {}
                
                translation_cache = {}
                # Define apply step whenever complex-table mode runs. If `_apply_docx_tables` were only defined
                # inside `if content_list`, requests with zero translatable cells would skip the definition but
                # still invoke `run_in_threadpool(_apply_docx_tables)` → UnboundLocalError.
                if is_complex_table_bool:
                    def _apply_docx_tables():
                        # === DIAGNOSTIC: dump DOCX table structure before any changes ===
                        print("[DOCX-DIAG] ==== Table Structure Diagnosis ====", flush=True)
                        for t_idx, table in enumerate(doc.tables):
                            tbl = table._tbl
                            tblPr = tbl.find(qn('w:tblPr'))
                            tblLayout_val = "NOT SET"
                            tblW_info = "NOT SET"
                            if tblPr is not None:
                                tblLayout_el = tblPr.find(qn('w:tblLayout'))
                                if tblLayout_el is not None:
                                    tblLayout_val = tblLayout_el.get(qn('w:type'))
                                tblW_el = tblPr.find(qn('w:tblW'))
                                if tblW_el is not None:
                                    tblW_info = f"{tblW_el.get(qn('w:w'))} ({tblW_el.get(qn('w:type'))})"
                            # tblGrid
                            try:
                                grid_cols = tbl.tblGrid.findall(qn('w:gridCol'))
                                grid_widths = [gc.get(qn('w:w')) for gc in grid_cols]
                            except Exception:
                                grid_widths = []
                            print(f"[DOCX-DIAG] Table {t_idx}: tblLayout={tblLayout_val} tblW={tblW_info} grid={grid_widths}", flush=True)
                            # First 2 rows
                            for r_idx, row in enumerate(table.rows[:2]):
                                for c_idx, cell in enumerate(row.cells[:4]):
                                    tc = cell._tc
                                    tcPr = tc.find(qn('w:tcPr'))
                                    tcW_info = noWrap_info = textDir_info = gridSpan_info = "N/A"
                                    if tcPr is not None:
                                        tcW_el = tcPr.find(qn('w:tcW'))
                                        tcW_info = f"{tcW_el.get(qn('w:w'))}({tcW_el.get(qn('w:type'))})" if tcW_el is not None else "None"
                                        noWrap_info = "YES" if tcPr.find(qn('w:noWrap')) is not None else "no"
                                        td_el = tcPr.find(qn('w:textDirection'))
                                        textDir_info = td_el.get(qn('w:val')) if td_el is not None else "None"
                                        gs_el = tcPr.find(qn('w:gridSpan'))
                                        gridSpan_info = gs_el.get(qn('w:val')) if gs_el is not None else "1"
                                    text_preview = "".join(p.text.strip() for p in cell.paragraphs if p.text.strip())[:20]
                                    print(f"[DOCX-DIAG]   r{r_idx}c{c_idx}: tcW={tcW_info} noWrap={noWrap_info} textDir={textDir_info} span={gridSpan_info} | '{text_preview}'", flush=True)
                        print("[DOCX-DIAG] ====================================", flush=True)
                        # === END DIAGNOSTIC ===

                        def apply_style(container, prefix=""):
                            if hasattr(container, 'tables'):
                                for t_idx, table in enumerate(container.tables):

                                    # === FIX: LibreOffice fixed-layout uses tblGrid, not tcW ===
                                    # Must sync tblGrid from actual cell widths AND set tblW=dxa.
                                    # For span=N cells, distribute width evenly across N gridCols.
                                    try:
                                        from docx.oxml import OxmlElement
                                        tblPr = table._tbl.find(qn('w:tblPr'))
                                        if tblPr is not None:
                                            # Build new tblGrid from actual TC elements (handles span correctly)
                                            tc_elements = table.rows[0]._tr.findall(qn('w:tc')) if table.rows else []
                                            new_grid_widths = {}
                                            grid_col_idx = 0
                                            total_w = 0
                                            for tc_el in tc_elements:
                                                tc_pr = tc_el.find(qn('w:tcPr'))
                                                width = 0
                                                span = 1
                                                if tc_pr is not None:
                                                    tcW_el = tc_pr.find(qn('w:tcW'))
                                                    if tcW_el is not None:
                                                        try:
                                                            width = int(tcW_el.get(qn('w:w')) or 0)
                                                        except Exception:
                                                            pass
                                                    span_el = tc_pr.find(qn('w:gridSpan'))
                                                    if span_el is not None:
                                                        try:
                                                            span = int(span_el.get(qn('w:val'), '1'))
                                                        except Exception:
                                                            span = 1
                                                total_w += width
                                                each = width // span if span > 0 else width
                                                for s in range(span):
                                                    new_grid_widths[grid_col_idx + s] = each
                                                # Fix rounding on last col of span
                                                if span > 1:
                                                    new_grid_widths[grid_col_idx + span - 1] += width - each * span
                                                grid_col_idx += span

                                            # Set tblW to explicit dxa
                                            if total_w > 0:
                                                tblW_el = tblPr.find(qn('w:tblW'))
                                                if tblW_el is None:
                                                    tblW_el = OxmlElement('w:tblW')
                                                    tblPr.insert(0, tblW_el)
                                                tblW_el.set(qn('w:w'), str(total_w))
                                                tblW_el.set(qn('w:type'), 'dxa')

                                            # Sync tblGrid cols to new widths
                                            tblGrid = table._tbl.tblGrid
                                            if tblGrid is not None and new_grid_widths:
                                                grid_cols = tblGrid.findall(qn('w:gridCol'))
                                                for i, gc in enumerate(grid_cols):
                                                    if i in new_grid_widths and new_grid_widths[i] > 0:
                                                        gc.set(qn('w:w'), str(new_grid_widths[i]))

                                            # Set tblLayout to fixed
                                            tblLayout = tblPr.find(qn('w:tblLayout'))
                                            if tblLayout is None:
                                                tblLayout = OxmlElement('w:tblLayout')
                                                tblPr.append(tblLayout)
                                            tblLayout.set(qn('w:type'), 'fixed')
                                    except Exception as e:
                                        print(f"[DOCX-WARN] tblLayout fix error: {e}", flush=True)




                                    for r_idx, row in enumerate(table.rows):
                                        # Allow rows to expand downward
                                        try:
                                            trPr = row._tr.trPr
                                            if trPr is not None:
                                                for trHeight in trPr.findall(qn('w:trHeight')):
                                                    trHeight.set(qn('w:hRule'), 'atLeast')
                                        except Exception:
                                            pass
                                            
                                        for c_idx, cell in enumerate(row.cells):
                                            key = f"{prefix}t{t_idx}_r{r_idx}_c{c_idx}"

                                            # --- Snapshot original cell width BEFORE changes ---
                                            original_tcW_w = None
                                            original_tcW_type = None
                                            try:
                                                tcPr = cell._tc.get_or_add_tcPr()
                                                tcW = tcPr.find(qn('w:tcW'))
                                                if tcW is not None:
                                                    original_tcW_w = tcW.get(qn('w:w'))
                                                    original_tcW_type = tcW.get(qn('w:type'))
                                            except Exception:
                                                pass

                                            # Remove noWrap to allow text wrapping
                                            try:
                                                tcPr = cell._tc.get_or_add_tcPr()
                                                for noWrap in tcPr.findall(qn('w:noWrap')):
                                                    tcPr.remove(noWrap)
                                            except Exception:
                                                pass

                                            original_text = cell_texts.get(key)
                                            translated_text = translation_cache.get(original_text) if original_text else None

                                            # Fallback: match live cell text directly
                                            if not translated_text and hasattr(cell, 'paragraphs'):
                                                live_text = "".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                                                translated_text = translation_cache.get(live_text)

                                            if translated_text and hasattr(cell, 'paragraphs'):
                                                paras = list(cell.paragraphs)
                                                if paras:
                                                    first_para = paras[0]
                                                    if first_para.runs:
                                                        first_para.runs[0].text = translated_text
                                                        for j in range(1, len(first_para.runs)):
                                                            first_para.runs[j].text = ""
                                                    else:
                                                        first_para.add_run(translated_text)
                                                    
                                                    for p_idx in range(1, len(paras)):
                                                        for run in paras[p_idx].runs:
                                                            run.text = ""
                                                    
                                                    first_para.paragraph_format.left_indent = None
                                                    first_para.paragraph_format.right_indent = None
                                                    first_para.paragraph_format.line_spacing = 1.2
                                                    
                                                    for run in first_para.runs:
                                                        if run.text:
                                                            try:
                                                                run.font.name = 'Arial'
                                                                run.font.size = Pt(11)
                                                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
                                                            except Exception:
                                                                pass

                                            # --- Restore original cell width AFTER changes ---
                                            # This prevents narrow category columns (e.g. 操作→Operation)
                                            # from expanding to fit the longer English text.
                                            if original_tcW_w is not None:
                                                try:
                                                    tcPr = cell._tc.get_or_add_tcPr()
                                                    tcW = tcPr.find(qn('w:tcW'))
                                                    if tcW is None:
                                                        from docx.oxml import OxmlElement
                                                        tcW = OxmlElement('w:tcW')
                                                        tcPr.append(tcW)
                                                    tcW.set(qn('w:w'), original_tcW_w)
                                                    if original_tcW_type:
                                                        tcW.set(qn('w:type'), original_tcW_type)
                                                except Exception:
                                                    pass

                                            # Recurse for nested tables
                                            apply_style(cell, f"{key}_")
                        
                        apply_style(doc)
                        doc.save(docx_path)

                if content_list:
                    print(f"[PDF-DOCX] Batch translating {len(content_list)} table cell strings...", flush=True)
                    batch_size = 50

                    async def process_batch(batch):
                        translated = await pdf_service._translate_batch_ollama(batch, target_lang=target_lang)
                        for orig, trans in zip(batch, translated):
                            if trans and "<SKIP>" not in trans:
                                translation_cache[orig] = trans

                    tasks = []
                    for i in range(0, len(content_list), batch_size):
                        tasks.append(process_batch(content_list[i : i + batch_size]))

                    await asyncio.gather(*tasks)

                if is_complex_table_bool:
                    # Complex mode: apply table translations then convert DOCX -> PDF (overwrites the layout-preserved PDF)
                    await run_in_threadpool(_apply_docx_tables)
                    print(f"[PDF-DOCX] Applied {len(translation_cache)} translations to DOCX tables.", flush=True)
                    
                    with open(docx_path, 'rb') as f:
                        docx_b64 = base64.b64encode(f.read()).decode('utf-8')

                    # Stage 3: Convert translated DOCX back to PDF
                    def _docx_to_pdf():
                        try:
                            if platform.system() == "Windows":
                                if docx2pdf_convert:
                                    docx2pdf_convert(docx_path, output_pdf_path)
                                else:
                                    print("[PDF-DOCX] docx2pdf library not found on Windows.", flush=True)
                            else:
                                subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', os.path.dirname(output_pdf_path), docx_path], check=True)
                        except Exception as e:
                            print(f"[PDF-DOCX] Docx to PDF failed: {e}", flush=True)

                    await run_in_threadpool(_docx_to_pdf)

                    # Reload the re-rendered PDF
                    if os.path.exists(output_pdf_path):
                        with open(output_pdf_path, 'rb') as f:
                            pdf_b64 = base64.b64encode(f.read()).decode('utf-8')
                else:
                    # Normal mode: DOCX is just the pdf2docx conversion of the already-translated PDF.
                    # No stage 2 table translation, no docx->pdf roundtrip (avoids layout corruption).
                    # The original layout-preserved PDF from stage 1 is already in pdf_b64.
                    print("[PDF-DOCX] Normal mode: skipping table re-translation and docx->pdf roundtrip.", flush=True)
                    with open(docx_path, 'rb') as f:
                        docx_b64 = base64.b64encode(f.read()).decode('utf-8')

                background_tasks.add_task(os.remove, docx_path)
            except Exception as docx_err:
                print(f"[PDF] Docx translation/conversion failed: {docx_err}")
                import traceback
                traceback.print_exc()
                if os.path.exists(docx_path): os.remove(docx_path)
                
        background_tasks.add_task(os.remove, temp_input_path)
        background_tasks.add_task(os.remove, output_pdf_path)
        
        return {"pdf_base64": pdf_b64, "docx_base64": docx_b64}
    except Exception as e:
        if temp_input_path and os.path.exists(temp_input_path): os.remove(temp_input_path)
        if docx_path and os.path.exists(docx_path): os.remove(docx_path)
        raise HTTPException(status_code=500, detail=str(e))

# ------------------------------------------------------------------------------
# n8n Microservice Interface: POST /api/v1/document/process
# Receives file + n8n-controlled parameters (model, temperature, target_lang).
# Delegates entirely to the existing /pdf-translation core logic.
# All file output / DOCX generation logic is untouched.
# ------------------------------------------------------------------------------
@app.post("/api/v1/document/process")
async def document_process_for_n8n(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    model: str = Form(""),                    # Ollama model override from n8n (empty = use service default)
    temperature: float = Form(0.1),           # Translation temperature from n8n
    target_lang: str = Form("zh-TW"),         # Target language forwarded from frontend via n8n
    is_complex_table: str = Form("false"),    # Forwarded from frontend toggle; true = YOLO protects tables, DOCX handles cells
    num_ctx: Optional[int] = Form(None),      # Context window size override from n8n (None = use service default)
):
    """
    n8n Document Translation Microservice Entrypoint.
    - Accepts model / temperature overrides from the n8n HTTP Request node.
    - target_lang is forwarded transparently from the frontend via the n8n Webhook.
    - Delegates to the existing /pdf-translation core logic without duplicating code.
    - Returns {"pdf_base64": ..., "docx_base64": ...} for n8n to forward back.
    """
    # Temporarily override pdf_service runtime settings
    original_model = pdf_service.ollama_model
    original_temperature = pdf_service.temperature

    model_override = model.strip() or None
    if model_override:
        pdf_service.ollama_model = model_override
    pdf_service.temperature = temperature
    
    # Store original and set new num_ctx if provided
    original_num_ctx = pdf_service.ollama_num_ctx
    if num_ctx:
        pdf_service.ollama_num_ctx = num_ctx

    print(
        f"[n8n Document] model={pdf_service.ollama_model} | "
        f"temperature={temperature} | target_lang={target_lang}",
        flush=True,
    )

    try:
        # Delegate to existing translate_pdf function (all core logic lives there)
        result = await translate_pdf(
            background_tasks=background_tasks,
            file=file,
            target_lang=target_lang,
            debug="false",
            is_complex_table=is_complex_table,
            num_ctx=num_ctx,
        )
        return result
    finally:
        # Always restore original settings to avoid affecting other concurrent requests
        pdf_service.ollama_model = original_model
        pdf_service.temperature = original_temperature
        pdf_service.ollama_num_ctx = original_num_ctx

@app.post("/chat")
async def chat_text(payload: dict):
    text = payload.get("question", payload.get("text", ""))
    with gpu_work_scope(MODULE_KM_ASSISTANT):
        response = await llm_service.chat(text)
    return {"response": response}

# ── n8n Sub-Agent Endpoints ────────────────────────────────────────────────────
# These endpoints are called by the two n8n branches (SQL_PROD / SQL_EQ).
# The Router decision is made in n8n; actual AI processing stays in the backend.

N8N_FACTORY_WEBHOOK = "http://172.16.2.68:5678/webhook/factory-chat"
N8N_DOC_WEBHOOK    = os.getenv("N8N_DOC_WEBHOOK",    "http://172.16.2.68:5678/webhook/doc-chat")
N8N_DOC_INGEST     = os.getenv("N8N_DOC_INGEST",     "http://172.16.2.68:5678/webhook/doc-ingest")

class FactoryChatRequest(BaseModel):
    question: str
    history: Optional[list] = []

@app.post("/api/v1/factory/sql/production-chat")
async def production_chat(request: FactoryChatRequest):
    """Production SQL Agent endpoint - called by n8n SQL_PROD branch"""
    try:
        print(f"[Production Chat] question='{request.question}'", flush=True)
        result = await factory_agent.sql_agent.chat(request.question, history=request.history)
        if isinstance(result, dict):
            return result
        return {"response": str(result), "chart_config": None}
    except Exception as e:
        print(f"[Production Chat Error] {e}", flush=True)
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/v1/factory/sql/equipment-chat")
async def equipment_chat(request: FactoryChatRequest):
    """Equipment SQL Agent endpoint - called by n8n SQL_EQ branch"""
    try:
        print(f"[Equipment Chat] question='{request.question}'", flush=True)
        result = await factory_agent.equipment_sql_agent.chat(request.question, history=request.history)
        if isinstance(result, dict):
            return result
        return {"response": str(result), "chart_config": None}
    except Exception as e:
        print(f"[Equipment Chat Error] {e}", flush=True)
        raise HTTPException(status_code=500, detail=str(e))

# ──────────────────────────────────────────────────────────────────────────────

@app.post("/factory-chat")
async def factory_chat(payload: dict, background_tasks: BackgroundTasks):
    """Factory data chat interface: proxies to n8n Router webhook.
    n8n decides the route (SQL_PROD / SQL_EQ) and calls the
    corresponding sub-agent endpoint on this server.
    """
    try:
        user_text = payload.get("text")
        session_id = payload.get("session_id")
        if not user_text: raise HTTPException(status_code=400, detail="Text field is required")
        
        history = []
        if session_id:
            session = await factory_store.get_session(session_id)
            if session: history = session.get("messages", [])
        
        # Ensure session id exists
        if not session_id:
            try:
                session = await factory_store.create_session(user_text)
                session_id = session["session_id"]
            except Exception as se:
                session_id = str(uuid.uuid4())
                print(f"[Session Warning] Fallback session id generated: {se}")

        print(f"\n[Factory Chat] Forwarding to n8n (session={session_id})", flush=True)

        # Forward to n8n Router webhook; n8n will call the sub-agent
        # endpoints above and return { response, chart_config }
        # Use longer read timeout to accommodate slow LLM + SQL round-trips via n8n
        async with httpx.AsyncClient(timeout=httpx.Timeout(connect=10.0, read=300.0, write=30.0, pool=10.0)) as client:
            n8n_resp = await client.post(
                N8N_FACTORY_WEBHOOK,
                json={"question": user_text, "history": history}
            )
            n8n_resp.raise_for_status()
            raw = n8n_resp.text.strip()
            if not raw:
                print("[Factory API Warning] n8n returned empty body – workflow may have an unhandled route.", flush=True)
                return {
                    "response": "抱歉，系統目前無法處理這個問題（後端工作流程未回傳任何資料）。請稍後再試，或換一個問法。",
                    "session_id": str(session_id),
                    "chart_config": None,
                }
            result = n8n_resp.json()

        response_text = result.get("response", "")
        chart_config  = result.get("chart_config", None)
        
        # Archive conversation history in background
        try:
            background_tasks.add_task(factory_store.append_messages, session_id, user_text, response_text)
        except Exception as bg_e:
            print(f"[BG Task Error] Failed to queue session save: {bg_e}")

        print(f"[Factory Chat] Success: {len(response_text)} chars", flush=True)
        
        return {
            "response": response_text,
            "session_id": str(session_id),
            "chart_config": chart_config
        }
        
    except Exception as e:
        print(f"[Factory API Error] {e}", flush=True)
        print(traceback.format_exc(), flush=True)
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/factory-sessions")
async def list_factory_sessions():
    try:
        sessions = await factory_store.list_sessions()
        return {"sessions": sessions}
    except Exception as e:
        print(f"[API Error] list_sessions: {e}")
        return {"sessions": []}

@app.get("/factory-sessions/{session_id}")
async def get_factory_session(session_id: str):
    session = await factory_store.get_session(session_id)
    if not session: raise HTTPException(status_code=404, detail="Session not found")
    return session

@app.delete("/factory-sessions/{session_id}")
async def delete_factory_session(session_id: str):
    deleted = await factory_store.delete_session(session_id)
    if not deleted: raise HTTPException(status_code=404, detail="Session not found")
    return {"message": "Session deleted"}

# ── Document Knowledge (RAG) Chat ──────────────────────────────────────────────
# Proxies to the n8n PDF RAG webhook.
# n8n receives {"question": "..."}, runs ChromaDB retrieval + LLM, and returns
# {"output": "answer text"} via a Respond to Webhook node.
@app.post("/document-chat")
async def document_chat(payload: dict, background_tasks: BackgroundTasks):
    """Document KM chat: proxies to n8n PDF RAG webhook with session tracking."""
    question = payload.get("question") or payload.get("text")
    session_id = payload.get("session_id")
    if not question:
        raise HTTPException(status_code=400, detail="question field is required")

    # Ensure session exists
    if not session_id:
        try:
            session = await doc_store.create_session(question)
            session_id = session["session_id"]
        except Exception as se:
            session_id = str(uuid.uuid4())
            print(f"[DocSession Warning] Fallback session id: {se}")

    try:
        print(f"\n[Document Chat] session={session_id} question='{question}'", flush=True)
        with gpu_work_scope(MODULE_KM_ASSISTANT):
            async with httpx.AsyncClient(
                timeout=httpx.Timeout(connect=10.0, read=180.0, write=30.0, pool=10.0)
            ) as client:
                n8n_resp = await client.post(
                    N8N_DOC_WEBHOOK,
                    json={"question": question},
                )
                n8n_resp.raise_for_status()
                raw = n8n_resp.text.strip()
                if not raw:
                    print("[Document Chat Warning] n8n returned empty body.", flush=True)
                    fallback = "抱歉，知識庫查詢未回傳結果，請稍後再試。"
                    background_tasks.add_task(doc_store.append_messages, session_id, question, fallback)
                    return {"response": fallback, "session_id": session_id}

                result = n8n_resp.json()

        # n8n returns {"output": "..."} or {"response": "..."}
        response_text = result.get("output") or result.get("response") or ""
        print(f"[Document Chat] Success: {len(response_text)} chars", flush=True)

        # Save message pair to session in background
        background_tasks.add_task(doc_store.append_messages, session_id, question, response_text)

        return {"response": response_text, "session_id": str(session_id)}

    except Exception as e:
        print(f"[Document Chat Error] {e}", flush=True)
        print(traceback.format_exc(), flush=True)
        # Save message pair even on error so the session is not empty
        err_msg = "⚠️ 知識庫服務暫時無法使用，請稍後再試。"
        background_tasks.add_task(doc_store.append_messages, session_id, question, err_msg)
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/document-sessions")
async def list_document_sessions():
    try:
        sessions = await doc_store.list_sessions()
        return {"sessions": sessions}
    except Exception as e:
        print(f"[API Error] list_document_sessions: {e}")
        return {"sessions": []}

@app.get("/document-sessions/{session_id}")
async def get_document_session(session_id: str):
    session = await doc_store.get_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    return session

@app.delete("/document-sessions/{session_id}")
async def delete_document_session(session_id: str):
    deleted = await doc_store.delete_session(session_id)
    if not deleted:
        raise HTTPException(status_code=404, detail="Session not found")
    return {"message": "Session deleted"}



# ── Document Ingest (PDF → ChromaDB via n8n) ───────────────────────────────────
# Accepts a PDF file upload and forwards it to the n8n doc-ingest webhook as
# multipart/form-data so n8n can extract text, chunk, embed and store in ChromaDB.
@app.post("/document-ingest")
async def document_ingest(
    file: UploadFile = File(...),
    session_id: str | None = Form(default=None),
):
    """Forward an uploaded PDF to the n8n PDF ingestion webhook.

    Optional `session_id` ties the file to a chat session for the UI file list.
    If omitted, a new empty session is created so uploads are always per-session.
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")
    # Only accept PDF and common document types
    allowed = {"application/pdf", "application/octet-stream"}
    if file.content_type not in allowed and not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Only PDF files are supported")

    # Temporarily throttle PDF translation to free up GPU compute for the embedding router.
    # We acquire extra slots from the translation semaphore so that at most
    # (total - INGEST_RESERVED) translation tasks can run while this upload is active.
    INGEST_RESERVED = int(os.getenv("INGEST_GPU_RESERVED_SLOTS", "10"))
    _throttle_holders: list[bool] = []
    sem = llm_service._translate_batch_sem
    for _ in range(INGEST_RESERVED):
        acquired = await asyncio.wait_for(asyncio.shield(sem.acquire()), timeout=2.0) if sem._value > 0 else None
        if acquired is not None:
            _throttle_holders.append(True)

    try:
        try:
            sid = (session_id or "").strip() or None
            if not sid:
                new_sess = await doc_store.create_empty_session("新對話")
                sid = new_sess["session_id"]
            else:
                if not await doc_store.get_session(sid):
                    raise HTTPException(status_code=404, detail="Session not found")

            print(f"[Document Ingest] Received: {file.filename} ({file.content_type}) session={sid}", flush=True)
            content = await file.read()
            async with httpx.AsyncClient(
                timeout=httpx.Timeout(connect=10.0, read=300.0, write=60.0, pool=10.0)
            ) as client:
                n8n_resp = await client.post(
                    N8N_DOC_INGEST,
                    files={"file": (file.filename, content, file.content_type or "application/pdf")}
                )
                n8n_resp.raise_for_status()
                raw = n8n_resp.text.strip()
                if not raw:
                    try:
                        await doc_store.add_file(sid, file.filename, len(content))
                    except Exception as fe:
                        print(f"[Document Ingest] Warning: failed to persist file metadata: {fe}", flush=True)
                    return {
                        "status": "ok",
                        "filename": file.filename,
                        "message": "Ingested (no detail returned)",
                        "session_id": sid,
                    }
                result = n8n_resp.json()

            print(f"[Document Ingest] Done: {file.filename}", flush=True)
            try:
                await doc_store.add_file(sid, file.filename, len(content))
            except Exception as fe:
                print(f"[Document Ingest] Warning: failed to persist file metadata: {fe}", flush=True)
            return {"status": "ok", "filename": file.filename, "session_id": sid, **result}

        except Exception as e:
            print(f"[Document Ingest Error] {e}", flush=True)
            print(traceback.format_exc(), flush=True)
            raise HTTPException(status_code=500, detail=str(e))
    finally:
        # Always release throttle slots back to the translation semaphore
        for _ in _throttle_holders:
            sem.release()



@app.get("/document-files")
async def list_document_files(session_id: str):
    """Return ingested file records for one chat session (UI list; same Chroma collection allowed)."""
    if not (session_id or "").strip():
        raise HTTPException(status_code=400, detail="session_id is required")
    files = await doc_store.list_files(session_id.strip())
    return {"files": files}


@app.delete("/document-files/{filename:path}")
async def delete_document_file(filename: str, session_id: str):
    """Remove a file record for a session from Redis. (Does not delete from ChromaDB.)"""
    if not (session_id or "").strip():
        raise HTTPException(status_code=400, detail="session_id is required")
    deleted = await doc_store.delete_file(session_id.strip(), filename)
    if not deleted:
        raise HTTPException(status_code=404, detail="File record not found")
    return {"message": f"{filename} removed from file index"}

